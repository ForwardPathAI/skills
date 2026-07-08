#!/usr/bin/env python3
"""Reusable DOCX generator for the poc-to-product-architecture deliverable.

This is a TEMPLATE. Copy it to `docs/architecture/build_docx.py` in the POC repo,
then edit the CONFIG block and the `build()` content region so it mirrors the
architecture canvas (`<app>-architecture.canvas.tsx`).

Hard rule (see DOCX_GUIDE.md): the DOCX embeds RENDERED diagram images only. It
must NOT contain literal Bicep or Mermaid source — those live in the canvas.
That is why this template ships no code-block helper.

Pipeline:
  1. Author mermaid sources next to this file: `<app>-architecture.mmd`,
     `<app>-cost.mmd`, and any extra diagram (e.g. `<app>-tenancy.mmd`).
  2. Render each to PNG:
       npx -y @mermaid-js/mermaid-cli -i <name>.mmd -o <name>.png -s 3 -b white
     (add `-p puppeteer.json` with {"args":["--no-sandbox","--disable-setuid-sandbox"]}
     if the renderer segfaults in a sandbox).
  3. Run this script:  python3 build_docx.py
     (needs `python-docx`: pip install python-docx)

Every table row can carry a semantic tone — "success" / "warning" / "info" /
"danger" — which tints the row and (optionally) a status cell, so the DOCX reads
like the color-coded canvas.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# ============================================================================
# CONFIG — edit these for the product being documented.
# ============================================================================
APP_SLUG = "my-app"                                    # image / diagram basename
APP_TITLE = "My App"                                   # cover title
APP_SUBTITLE = "Production Architecture — from POC to product"
# Footer provenance: cite the SOW version/date so the doc is traceable.
FOOTER_NOTE = f"{APP_TITLE} — Production Architecture  ·  Grounded in SOW  ·  Forward Path"
OUTFILE = f"{APP_TITLE.replace(' ', '-')}-Production-Architecture.docx"

# Diagram PNGs expected next to this script (render from the sibling .mmd files).
ARCH_PNG = f"{APP_SLUG}-architecture.png"              # required: target topology
COST_PNG = f"{APP_SLUG}-cost.png"                      # optional: cost chart (xychart-beta)
# Add extra diagrams here as (png_filename, width_inches, caption) and place them
# wherever they belong in build() with add_image(...).

# ----- palette (Forward Path deliverable house style) ---------------------
NAVY = "1F3B63"        # header row / title ink
BLUE = "2E79B5"        # section headings
INK = "0B1E3F"
GRAY_TEXT = "5B6472"
BORDER = "D6DCE4"
STRIPE = "F5F7FA"

ROW_TINT = {"success": "F1F9F5", "warning": "FEF8EF", "info": "F2F7FD", "danger": "FDF1EF"}
CELL_TINT = {"success": "CDEBD9", "warning": "FBE4C0", "info": "D4E4F7", "danger": "F6D2CC"}
CALLOUT = {"info": ("EAF1FB", "2E79B5"), "warning": ("FDF3E2", "C8871B")}


# ============================================================================
# Low-level XML helpers (proven — usually no need to edit below this line).
# ============================================================================
def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_border(tcPr, edge, color, sz="4"):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def cell_borders(cell, color=BORDER):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        borders.append(_set_border(tcPr, edge, color))
    tcPr.append(borders)


def cell_margins(cell, top=60, bottom=60, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def set_run(run, size=10.5, bold=False, italic=False, color=None, font="Calibri", mono=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    name = "Consolas" if mono else font
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def fill_cell(cell, segments, align="left", size=9.5, bg=None):
    """segments: str or list of (text, opts-dict). opts: bold/italic/color/size/mono."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if isinstance(segments, str):
        segments = [(segments, {})]
    for text, opts in segments:
        r = p.add_run(text)
        set_run(r, size=opts.get("size", size), bold=opts.get("bold", False),
                italic=opts.get("italic", False), color=opts.get("color", None),
                mono=opts.get("mono", False))
    if bg:
        shade(cell, bg)
    cell_borders(cell)
    cell_margins(cell)
    vcenter(cell)


# ============================================================================
# Building blocks
# ============================================================================
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_run(r, size=16, bold=True, color=BLUE)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), BORDER)
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        set_run(r, size=12.5, bold=True, color=INK)
    return p


def add_body(doc, segments, size=10.5, space_after=6, align="left"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    if isinstance(segments, str):
        segments = [(segments, {})]
    for text, opts in segments:
        r = p.add_run(text)
        set_run(r, size=opts.get("size", size), bold=opts.get("bold", False),
                italic=opts.get("italic", False), color=opts.get("color", GRAY_TEXT),
                mono=opts.get("mono", False))
    return p


def add_bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        segs = it if isinstance(it, list) else [(it, {})]
        for text, opts in segs:
            r = p.add_run(text)
            set_run(r, size=opts.get("size", size), bold=opts.get("bold", False),
                    color=opts.get("color", "2B2B2B"))


def fixed_layout(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def add_table(doc, headers, rows, widths, aligns=None, row_tones=None,
              tone_col=None, header_bg=NAVY, striped=True, body_size=9.5):
    """rows: list of rows; each cell is a str OR a list of (text, opts) segments.

    row_tones[i] in {success,warning,info,danger} tints row i; if tone_col is set,
    that cell gets a stronger tint — ideal for a Status / Tag / Severity column.
    """
    aligns = aligns or ["left"] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fixed_layout(table)
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        fill_cell(c, [(h, {"bold": True, "color": "FFFFFF", "size": 9.5})],
                  align=aligns[j], bg=header_bg)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        tone = row_tones[i] if row_tones else None
        base_bg = ROW_TINT[tone] if tone else (STRIPE if (striped and i % 2 == 1) else None)
        for j, val in enumerate(row):
            bg = base_bg
            if tone_col is not None and j == tone_col and tone:
                bg = CELL_TINT[tone]
            segs = val if isinstance(val, list) else [(str(val), {})]
            extra = {"bold": True, "color": "333333"} if (tone_col is not None and j == tone_col) else {}
            if extra:
                segs = [(t, {**o, **extra}) for t, o in segs]
            fill_cell(cells[j], segs, align=aligns[j], size=body_size, bg=bg)
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, tone, title, body_lines):
    """tone in {info, warning}. body_lines: list of segment-lists."""
    bg, bar = CALLOUT[tone]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fixed_layout(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.1)
    cell.text = ""
    shade(cell, bg)
    cell_borders(cell, bar)
    cell_margins(cell, top=110, bottom=110, left=150, right=140)
    pt = cell.paragraphs[0]
    pt.paragraph_format.space_after = Pt(4)
    rt = pt.add_run(title)
    set_run(rt, size=10.5, bold=True, color=bar)
    for segs in body_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        for text, opts in segs:
            r = p.add_run(text)
            set_run(r, size=9.5, bold=opts.get("bold", False), color=opts.get("color", "2B2B2B"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, filename, width_in, caption=None):
    """Embed a rendered diagram PNG. Skips gracefully if the file is missing so
    the doc still builds while you iterate on diagrams."""
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        warn = doc.add_paragraph()
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = warn.add_run(f"[missing diagram: {filename} — render the .mmd to PNG]")
        set_run(r, size=9, italic=True, color="B23B3B")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        r = c.add_run(caption)
        set_run(r, size=8.5, italic=True, color=GRAY_TEXT)


def add_stat_cards(doc, stats):
    """stats: list of (value, label) — a compact 'at a glance' band under the intro."""
    n = len(stats)
    st = doc.add_table(rows=2, cols=n)
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    st.autofit = False
    fixed_layout(st)
    for j, (val, lab) in enumerate(stats):
        fill_cell(st.rows[0].cells[j], [(val, {"bold": True, "size": 17, "color": BLUE})], align="center", bg="F2F7FD")
        fill_cell(st.rows[1].cells[j], [(lab, {"size": 8.5, "color": GRAY_TEXT})], align="center", bg="F2F7FD")
    usable = 7.1 / n
    for row in st.rows:
        for j in range(n):
            row.cells[j].width = Inches(usable)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_document():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("2B2B2B")

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run(FOOTER_NOTE), size=8, color="9AA3AE")

    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    set_run(tp.add_run(APP_TITLE), size=26, bold=True, color=INK)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    set_run(sp.add_run(APP_SUBTITLE), size=13, color=BLUE)
    return doc


# ============================================================================
# CONTENT — edit everything below to mirror the architecture canvas.
# Section order matches CANVAS_GUIDE.md. The Bicep skeleton (canvas section 6)
# and any literal mermaid source are intentionally OMITTED here — the DOCX
# shows the rendered architecture (and cost) diagrams instead.
# ============================================================================
def build():
    doc = setup_document()

    # 1. Executive overview -------------------------------------------------
    add_body(doc, [
        ("TODO one-paragraph product + architecture summary. Name the target stack "
         "(product-foundation) and deployment model (Container Apps + single Bicep, "
         "webhook install into the customer tenant). End by grounding it in the SOW: ", {}),
        ("SOW <vendor> <date> <title> (v1.0)", {"bold": True, "color": "2B2B2B"}),
        (".", {}),
    ], align="justify")

    add_stat_cards(doc, [
        ("$X–Y", "Azure run-rate / mo"),
        ("N+", "users"),
        ("W wk", "build timeline"),
        ("Region", "data residency"),
    ])

    add_callout(doc, "info", "Design intent vs. SOW", [[
        ("TODO one or two sentences: this maps the SOW onto the Forward Path "
         "product-foundation stack and the Bicep / customer-tenant deployability "
         "contract; deviations are called out in Architecture and Risks.", {})
    ]])

    # 2. Architecture (diagram + layers + deviations) -----------------------
    add_heading(doc, "Architecture")
    add_image(doc, ARCH_PNG, 7.0,
              "Figure 1 — Target topology. Solid = sync request / data · dashed = async / external / secrets. "
              "Blue = compute (Container Apps) · green = data plane · purple = external / identity / AI · gray = client.")

    add_heading(doc, "Layers (product-foundation mapping)", level=2)
    # TODO: one row per applicable layer. col2 is the production path (mono).
    layers = [
        ["Web", "apps/web", "Next.js App Router dashboards; TanStack Query"],
        ["API", "apps/api", "Hono RPC + zValidator"],
        ["Data", "packages/database", "Drizzle + PostgreSQL"],
        ["Jobs", "apps/jobs", "BullMQ + Redis workers (if needed)"],
        ["Auth", "packages/auth", "Better Auth + Microsoft Entra SSO + org RBAC"],
        ["Shared", "packages/shared", "Zod schemas, route constants, types"],
    ]
    rows = [[[(r[0], {"bold": True, "color": NAVY})], [(r[1], {"mono": True, "size": 9})], r[2]] for r in layers]
    add_table(doc, ["Layer", "Production path", "Notes"], rows, [1.0, 1.9, 4.2],
              ["left", "left", "left"])

    add_callout(doc, "warning", "Deviations from SOW / product-foundation (confirm with customer)", [
        [("TODO deviation title. ", {"bold": True}),
         ("Why the design departs from the SOW-named infra, and the in-tenant equivalent.", {})],
    ])

    # 3. SOW traceability ---------------------------------------------------
    add_heading(doc, "SOW traceability")
    add_body(doc, "Every SOW requirement appears once. Green = in the POC, amber = partial, blue = net-new.",
             size=9, space_after=4)
    # TODO: (requirement, production component, status label, tone)
    sow = [
        ("TODO requirement (§x)", "packages/... component", "new", "info"),
        ("TODO requirement already in POC", "existing component", "covered by POC", "success"),
        ("TODO requirement partially scaffolded", "component", "partial", "warning"),
    ]
    rows = [[r[0], [(r[1], {"mono": True, "size": 8.5})], r[2]] for r in sow]
    add_table(doc, ["SOW requirement", "Production component", "Status"], rows,
              [3.4, 2.4, 1.3], ["left", "left", "center"],
              row_tones=[r[3] for r in sow], tone_col=2, striped=False)

    # 4. POC gap audit (omit entirely if there is no POC) -------------------
    add_heading(doc, "POC gap audit")
    add_body(doc, "POC repo: TODO owner/repo, audited against the SOW. "
                  "Red = must be replaced for production, amber = no trace yet, green = reusable as-is.",
             size=9, space_after=4)
    # TODO: (finding, tag, citation, tone). tags: demo-grade / missing / reusable
    poc = [
        ("TODO demo-grade finding", "demo-grade", "path/to/file.ts", "danger"),
        ("TODO missing capability", "missing", "path/or/§clause", "warning"),
        ("TODO reusable asset", "reusable", "path/to/file.ts", "success"),
    ]
    rows = [[r[0], r[1], [(r[2], {"mono": True, "size": 8})]] for r in poc]
    add_table(doc, ["Finding vs. SOW", "Tag", "Citation"], rows,
              [4.0, 1.1, 2.0], ["left", "center", "left"],
              row_tones=[r[3] for r in poc], tone_col=1, striped=False)

    # 5. Azure resource map -------------------------------------------------
    add_heading(doc, "Azure resource map")
    add_body(doc, [("All resources in the chosen region (state residency if the SOW requires it). Images: "
                    f"forwardpathai.azurecr.io/{APP_SLUG}-{{web,api,jobs}} — tag ", {}),
                   ("dev", {"bold": True, "color": "2B2B2B"}),
                   (" on main, ", {}),
                   ("semver + latest", {"bold": True, "color": "2B2B2B"}),
                   (" on GitHub Release.", {})], size=9, space_after=4)
    # TODO: (resource, sku/tier, region, purpose)
    azure = [
        ["Log Analytics + App Insights", "Pay-as-you-go", "Region", "Telemetry"],
        ["Container Apps environment", "Consumption", "Region", "Shared compute"],
        ["Container App — web", "0.5 vCPU / 1 GiB", "Region", "Next.js (external)"],
        ["Container App — api", "0.5 vCPU / 1 GiB", "Region", "Hono RPC (external)"],
        ["PostgreSQL Flexible Server", "B1ms dev / GP prod", "Region", "App data + auth · PITR"],
        ["Key Vault", "Standard", "Region", "Secret refs (names only)"],
    ]
    rows = [[[(r[0], {"bold": True, "color": NAVY, "size": 9})], r[1], r[2], r[3]] for r in azure]
    add_table(doc, ["Resource", "SKU / tier", "Region", "Purpose"], rows,
              [2.2, 1.7, 1.15, 2.05], ["left", "left", "left", "left"])

    # 6. Security and reliability (decisions, not aspirations) --------------
    add_heading(doc, "Security and reliability")
    add_heading(doc, "Security decisions", level=2)
    add_bullets(doc, [
        "TODO: secrets are Key Vault references only — names, not values.",
        "TODO: each Container App has a managed identity with AcrPull on forwardpathai; narrow RBAC.",
        "TODO: Zod validation at every API boundary; Better Auth + Entra SSO with org RBAC.",
    ])
    add_heading(doc, "Reliability decisions", level=2)
    add_bullets(doc, [
        "TODO: liveness + readiness probes; internal-only ingress for workers.",
        "TODO: min replicas >= 1 for user-facing services; scale-to-zero for jobs.",
        "TODO: Postgres PITR + backups; Log Analytics wired; rollback via image tags.",
    ])

    # 7. Cost estimate (table + optional rendered chart) --------------------
    add_heading(doc, "Cost estimate")
    add_body(doc, "TODO scale assumptions and what's in / out of scope. If the SOW gave no budget "
                  "signals, default to consumption / scale-to-zero SKUs and say so here.", size=9.5)
    cost_rows, cost_meta = [], []

    def crow(label, low, high, note, meta=None, bold=False):
        seg_l = [(label, {"bold": bold, "color": NAVY if meta == "group" else "2B2B2B"})]
        cost_rows.append([seg_l, [(low, {"bold": bold})], [(high, {"bold": bold})], [(note, {"size": 9})]])
        cost_meta.append(meta)

    # TODO: fill real line items. meta: group (navy) / sub (light) / total (blue) / None
    crow("A · Azure infrastructure", "", "", "", meta="group", bold=True)
    crow("Container Apps", "$X", "$Y", "consolidated; scale-to-zero")
    crow("PostgreSQL Flexible", "$X", "$Y", "PITR")
    crow("Subtotal — infra", "$X", "$Y", "", meta="sub", bold=True)
    crow("Azure cloud run-rate / mo", "~$X", "~$Y", "vs SOW budget signal", meta="total", bold=True)

    tbl = add_table(doc, ["Line item", "Low / mo", "High / mo", "Notes"], cost_rows,
                    [2.75, 1.05, 1.15, 2.15], ["left", "right", "right", "left"], striped=False)
    for i, m in enumerate(cost_meta):
        if m is None:
            continue
        bg = {"group": NAVY, "sub": "EEF3F9", "total": "D4E4F7"}[m]
        for c in tbl.rows[i + 1].cells:
            shade(c, bg)
        if m == "group":
            for c in tbl.rows[i + 1].cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor.from_string("FFFFFF")

    add_image(doc, COST_PNG, 5.8,
              "Figure 2 — Monthly Azure run-rate the architecture drives, against the SOW's own cost signal.")

    # 8. Phased migration path ---------------------------------------------
    add_heading(doc, "Phased migration path")
    # TODO: (phase, built/migrated, POC items replaced) aligned to SOW phasing.
    phases = [
        ["Phase 1", "TODO what gets built", "TODO POC items replaced"],
    ]
    rows = [[[(r[0], {"bold": True, "color": NAVY, "size": 9})], r[1], r[2]] for r in phases]
    add_table(doc, ["Phase", "Built / migrated", "POC items replaced"], rows,
              [1.6, 3.2, 2.3], ["left", "left", "left"])

    # 9. Risks and open questions ------------------------------------------
    add_heading(doc, "Risks and open questions")
    # TODO: (risk, severity, mitigation, tone)
    risks = [
        ("TODO Bicep-contract conflict or open question", "medium", "Proposed alternative / next action", "warning"),
        ("TODO inferred item awaiting confirmation", "low", "Who confirms and when", "info"),
    ]
    rows = [[r[0], r[1], r[2]] for r in risks]
    add_table(doc, ["Risk / open question", "Severity", "Mitigation"], rows,
              [3.55, 1.0, 2.55], ["left", "center", "left"],
              row_tones=[r[3] for r in risks], tone_col=1, striped=False)

    add_heading(doc, "Next-step skills", level=2)
    add_body(doc, [
        ("azure-infra-setup", {"bold": True, "color": "2B2B2B"}),
        (" — full Bicep / Terraform authoring and CI/CD wiring.   ", {}),
        ("customer-deployment-package", {"bold": True, "color": "2B2B2B"}),
        (" — customer-facing external Bicep variant and deployment handoff.", {}),
    ], size=10)

    out = os.path.join(HERE, OUTFILE)
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    build()
